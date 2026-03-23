import streamlit as st
import PyPDF2
from openai import OpenAI
from docx import Document
from io import BytesIO

# 1. Configuración de la página
st.set_page_config(page_title="Asistente Legal", page_icon="⚖️")

# 2. Master Prompt
PROMPT_MAESTRO = """Actúa como un Abogado Especialista.
Analiza el texto que se enviará a continuación. 
Espera a leer todas las partes para generar un informe final con:
- Índice Estructural
- Matriz de Defensa
- Tabla de Jurisprudencia Invocada
- Alertas Procesales.
Usa un tono profesional y técnico."""

# Función mejorada para crear el archivo Word con Tablas
def crear_word(texto_ia):
    doc = Document()
    doc.add_heading('Reporte de Análisis Legal', 0)
    
    lineas = texto_ia.split('\n')
    en_tabla = False
    tabla_word = None

    for linea in lineas:
        linea_limpia = linea.strip()
        
        # Ignorar líneas vacías
        if not linea_limpia:
            continue
            
        # Detectar si es una línea de tabla de Markdown (empieza y termina con |)
        if linea_limpia.startswith('|') and linea_limpia.endswith('|'):
            # Ignorar la línea separadora de Markdown (|---|---|)
            if set(linea_limpia.replace('|', '').replace('-', '').replace(' ', '')) == set():
                continue
                
            # Extraer los textos de cada celda
            celdas = [celda.strip() for celda in linea_limpia.split('|') if celda.strip() != '']
            
            if not en_tabla:
                # Es la primera fila, creamos la tabla en Word
                tabla_word = doc.add_table(rows=1, cols=len(celdas))
                tabla_word.style = 'Table Grid' # Le da el formato de cuadrícula visible
                hdr_cells = tabla_word.rows[0].cells
                for i, valor in enumerate(celdas):
                    hdr_cells[i].text = valor.replace('**', '') # Quitamos los asteriscos de negrita
                en_tabla = True
            else:
                # Es una fila de datos, la agregamos a la tabla existente
                row_cells = tabla_word.add_row().cells
                for i, valor in enumerate(celdas):
                    if i < len(row_cells):
                        row_cells[i].text = valor.replace('**', '')
        else:
            en_tabla = False # Ya no estamos en una tabla
            
            # Detectar Títulos
            if linea_limpia.startswith('###'):
                doc.add_heading(linea_limpia.replace('###', '').strip(), level=2)
            elif linea_limpia.startswith('##'):
                doc.add_heading(linea_limpia.replace('##', '').strip(), level=1)
            elif linea_limpia.startswith('#'):
                doc.add_heading(linea_limpia.replace('#', '').strip(), level=1)
            else:
                # Párrafo normal (le quitamos los asteriscos sueltos para que quede limpio)
                doc.add_paragraph(linea_limpia.replace('**', ''))
                
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# 3. Interfaz de Usuario
st.title("⚖️ Analizador de Demandas")
st.sidebar.header("⚙️ Configuración")
mi_llave = st.sidebar.text_input("Tu API Key de OpenAI", type="password")

archivo_subido = st.file_uploader("Sube el PDF del expediente", type="pdf")

if st.button("🚀 Iniciar Análisis"):
    if not mi_llave:
        st.warning("Falta la API Key en el panel izquierdo.")
    elif archivo_subido is None:
        st.warning("Por favor, sube un archivo PDF.")
    else:
        cliente = OpenAI(api_key=mi_llave)
        texto_completo = ""
        
        with st.spinner('Procesando documento...'):
            lector = PyPDF2.PdfReader(archivo_subido)
            for pagina in lector.pages:
                texto_completo += pagina.extract_text() + "\n"
            
            try:
                respuesta = cliente.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[
                        {"role": "system", "content": PROMPT_MAESTRO},
                        {"role": "user", "content": texto_completo}
                    ],
                    temperature=0.2
                )
                
                resultado = respuesta.choices[0].message.content
                st.success("¡Análisis Terminado!")
                st.markdown(resultado)
                
                # Generar el archivo Word en memoria con las nuevas tablas
                archivo_word = crear_word(resultado)
                
                # Botón de descarga de Word
                st.download_button(
                    label="📥 Descargar Reporte en Word",
                    data=archivo_word,
                    file_name="Reporte_Legal_Formateado.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            except Exception as e:
                st.error(f"Error al conectar con la IA: {e}")
